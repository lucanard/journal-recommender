"""
Report Generator — Creates a downloadable Word document from recommendation results.
Uses python-docx. Install: pip install python-docx
"""

import io
import logging
from datetime import datetime

log = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log.warning("python-docx not installed. Report generation disabled. Install with: pip install python-docx")


def _set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell. Each border is a dict with sz, color, val."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn('w:tcBorders'), {})
    for edge, props in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if props:
            elem = borders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): props.get('val', 'single'),
                qn('w:sz'): str(props.get('sz', 4)),
                qn('w:color'): props.get('color', 'CCCCCC'),
                qn('w:space'): '0',
            })
            borders.append(elem)
    tcPr.append(borders)


def _add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    return run


def generate_report(results: dict, abstract: str, constraints: dict = None) -> io.BytesIO:
    """
    Generate a Word document summarizing the recommendation results.

    Args:
        results: The full API response from /recommend
        abstract: The user's original abstract
        constraints: Optional dict of constraints used

    Returns:
        BytesIO buffer containing the .docx file
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x2C, 0x24, 0x17)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ('Heading 1', 20, (0x0B, 0x11, 0x20)),
        ('Heading 2', 14, (0x16, 0x20, 0x36)),
        ('Heading 3', 12, (0x8B, 0x69, 0x41)),
    ]:
        h = doc.styles[level]
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(*color)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(8)

    # ── Colors ──
    GOLD = (0xD4, 0xA0, 0x41)
    DARK = (0x0B, 0x11, 0x20)
    MID = (0x5C, 0x50, 0x40)
    MUTED = (0x93, 0x8A, 0x79)
    GREEN = (0x3E, 0xBF, 0x6F)

    # ════════════════════════════════════════════════════════════════
    #  TITLE BLOCK
    # ════════════════════════════════════════════════════════════════

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    _add_run(title_p, "Journal Recommendation Report", bold=True, size=22, color=DARK)

    # Subtitle / timestamp
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    now = datetime.now().strftime("%B %d, %Y at %H:%M")
    _add_run(sub_p, f"Generated on {now}", size=9.5, color=MUTED)
    _add_run(sub_p, "  ·  ", size=9.5, color=MUTED)
    _add_run(sub_p, "AI Journal Recommender v1.1", size=9.5, color=MUTED)

    # Divider
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(12)
    pBdr = div_p._element.get_or_add_pPr().makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:color'): 'D4A041', qn('w:space'): '1'
    })
    pBdr.append(bottom)
    div_p._element.get_or_add_pPr().append(pBdr)

    # ════════════════════════════════════════════════════════════════
    #  ANALYSIS SUMMARY
    # ════════════════════════════════════════════════════════════════

    if results.get("analysis_summary"):
        doc.add_heading("Analysis Summary", level=1)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        _add_run(p, results["analysis_summary"], size=10.5, color=MID)

    # ════════════════════════════════════════════════════════════════
    #  SEARCH PARAMETERS
    # ════════════════════════════════════════════════════════════════

    doc.add_heading("Search Parameters", level=1)

    # Abstract (truncated for report)
    doc.add_heading("Abstract", level=3)
    abstract_display = abstract[:800] + ("..." if len(abstract) > 800 else "")
    p = doc.add_paragraph()
    _add_run(p, abstract_display, italic=True, size=10, color=MID)

    # Constraints table
    if constraints:
        doc.add_heading("Constraints Applied", level=3)
        constraint_rows = []

        if constraints.get("discipline") and constraints["discipline"] != "Any":
            constraint_rows.append(("Discipline", constraints["discipline"]))
        if constraints.get("article_type"):
            constraint_rows.append(("Article Type", constraints["article_type"]))
        if constraints.get("oa_preference") and constraints["oa_preference"] != "Any":
            constraint_rows.append(("Open Access", constraints["oa_preference"]))
        if constraints.get("apc_free_only"):
            constraint_rows.append(("APC", "Free to publish only"))
        elif constraints.get("max_apc"):
            constraint_rows.append(("Max APC", f"${constraints['max_apc']:,.0f}"))
        if constraints.get("min_impact_factor"):
            constraint_rows.append(("Min Impact Factor", str(constraints["min_impact_factor"])))
        if constraints.get("target_impact") and constraints["target_impact"] != "Any":
            constraint_rows.append(("Target Impact", constraints["target_impact"]))
        if constraints.get("indexing_required"):
            constraint_rows.append(("Required Indexing", ", ".join(constraints["indexing_required"])))
        if constraints.get("keywords"):
            constraint_rows.append(("Keywords", constraints["keywords"]))

        if constraint_rows:
            table = doc.add_table(rows=len(constraint_rows), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = True

            for i, (label, value) in enumerate(constraint_rows):
                row = table.rows[i]

                # Label cell
                cell_l = row.cells[0]
                cell_l.width = Cm(4.5)
                p = cell_l.paragraphs[0]
                _add_run(p, label, bold=True, size=10, color=DARK)
                _set_cell_shading(cell_l, 'F5F3EE')
                border_def = {'sz': 4, 'color': 'E8E4DB'}
                _set_cell_border(cell_l, bottom=border_def)

                # Value cell
                cell_v = row.cells[1]
                p = cell_v.paragraphs[0]
                _add_run(p, value, size=10, color=MID)
                _set_cell_border(cell_v, bottom=border_def)

    # Pipeline stats
    stats_p = doc.add_paragraph()
    stats_p.paragraph_format.space_before = Pt(10)
    _add_run(stats_p, "Pipeline: ", bold=True, size=9.5, color=MUTED)
    searched = results.get("candidates_searched", "?")
    filtered = results.get("candidates_after_filter", "?")
    shown = len(results.get("recommendations", []))
    _add_run(stats_p, f"{searched} candidates searched → {filtered} after filtering → {shown} shown", size=9.5, color=MUTED)

    if results.get("constraints_relaxed"):
        warn_p = doc.add_paragraph()
        _add_run(warn_p, "⚠ ", size=10)
        _add_run(warn_p, "No journals matched all constraints. Results shown without filters.", bold=True, size=10, color=(0xA6, 0x7B, 0x28))

    # ════════════════════════════════════════════════════════════════
    #  RECOMMENDATIONS
    # ════════════════════════════════════════════════════════════════

    doc.add_heading("Recommended Journals", level=1)

    recs = results.get("recommendations", [])
    for rec in recs:
        rank = rec.get("rank", "?")
        name = rec.get("journal_name", "Unknown Journal")
        publisher = rec.get("publisher", "")
        fit = rec.get("fit", "Medium")

        # ── Journal heading ──
        h = doc.add_heading(level=2)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(4)
        _add_run(h, f"#{rank}  ", bold=True, size=11, color=GOLD)
        _add_run(h, name, bold=True, size=14, color=DARK)
        fit_color = GREEN if fit == "High" else GOLD if fit == "Medium" else (0xE8, 0x5D, 0x5D)
        _add_run(h, f"   [{fit} Fit]", bold=True, size=10, color=fit_color)

        # ── Publisher ──
        pub_p = doc.add_paragraph()
        pub_p.paragraph_format.space_after = Pt(2)
        _add_run(pub_p, publisher, size=10, color=MUTED)
        if rec.get("homepage"):
            _add_run(pub_p, f"  ·  {rec['homepage']}", size=9, color=GOLD)

        # ── Key metrics table ──
        metrics = []
        if rec.get("oa_model"):
            metrics.append(("Publishing Model", rec["oa_model"]))
        apc = rec.get("apc_estimate", "")
        if apc and apc not in ("?", "N/A"):
            metrics.append(("APC", apc))
        if rec.get("impact_factor") and rec["impact_factor"] not in ("?", "Unknown", ""):
            metrics.append(("Impact Factor", str(rec["impact_factor"])))
        if rec.get("acceptance_rate") and rec["acceptance_rate"] not in ("N/A", "Not publicly available"):
            metrics.append(("Acceptance Rate", rec["acceptance_rate"]))
        if rec.get("review_time") and rec["review_time"] not in ("N/A", "Not publicly available"):
            metrics.append(("Review Time", rec["review_time"]))
        if rec.get("indexing"):
            metrics.append(("Indexing", ", ".join(rec["indexing"])))
        if rec.get("impact_proxy") and rec["impact_proxy"] != "Unknown":
            metrics.append(("Impact Tier", rec["impact_proxy"]))

        if metrics:
            table = doc.add_table(rows=len(metrics), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for i, (label, value) in enumerate(metrics):
                row = table.rows[i]
                cell_l = row.cells[0]
                cell_l.width = Cm(4.0)
                p = cell_l.paragraphs[0]
                _add_run(p, label, bold=True, size=9.5, color=DARK)
                _set_cell_shading(cell_l, 'F5F3EE')
                b = {'sz': 4, 'color': 'E8E4DB'}
                _set_cell_border(cell_l, bottom=b)

                cell_v = row.cells[1]
                p = cell_v.paragraphs[0]
                _add_run(p, value, size=9.5, color=MID)
                _set_cell_border(cell_v, bottom=b)

        # ── Reasons ──
        reasons = rec.get("reasons", [])
        if reasons:
            reasons_h = doc.add_paragraph()
            reasons_h.paragraph_format.space_before = Pt(10)
            reasons_h.paragraph_format.space_after = Pt(4)
            _add_run(reasons_h, "WHY IT MATCHES", bold=True, size=9, color=GOLD, font_name='Calibri')

            for reason in reasons:
                rp = doc.add_paragraph()
                rp.paragraph_format.space_after = Pt(2)
                rp.paragraph_format.left_indent = Cm(0.5)
                _add_run(rp, "✓  ", bold=True, size=10, color=GREEN)
                _add_run(rp, reason, size=10, color=MID)

        # ── Concern ──
        concern = rec.get("concern", "")
        if concern:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_before = Pt(6)
            cp.paragraph_format.left_indent = Cm(0.5)
            _add_run(cp, "Note: ", bold=True, size=10, color=(0xA6, 0x7B, 0x28))
            _add_run(cp, concern, size=10, color=(0xA6, 0x7B, 0x28))

        # ── Subjects ──
        subjects = rec.get("subjects", [])
        if subjects:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            _add_run(sp, "Subjects: ", bold=True, size=9, color=MUTED)
            _add_run(sp, "  ·  ".join(subjects), size=9, color=MUTED)

        # Similarity score
        score = rec.get("score", 0)
        if score > 0:
            sc_p = doc.add_paragraph()
            sc_p.paragraph_format.space_after = Pt(8)
            _add_run(sc_p, f"Semantic similarity: {score:.4f}", size=9, color=MUTED)

    # ════════════════════════════════════════════════════════════════
    #  TIMING & FOOTER
    # ════════════════════════════════════════════════════════════════

    timing = results.get("timing", {})
    if timing:
        doc.add_paragraph()  # spacer
        div2 = doc.add_paragraph()
        pBdr2 = div2._element.get_or_add_pPr().makeelement(qn('w:pBdr'), {})
        top_b = pBdr2.makeelement(qn('w:top'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:color'): 'E8E4DB', qn('w:space'): '1'
        })
        pBdr2.append(top_b)
        div2._element.get_or_add_pPr().append(pBdr2)

        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = []
        for key, label in [("embed_ms", "embed"), ("search_ms", "search"), ("filter_ms", "filter"), ("rerank_ms", "rerank")]:
            if key in timing:
                parts.append(f"{label}: {timing[key]}ms")
        _add_run(tp, "  ·  ".join(parts), size=8.5, color=MUTED, font_name='Calibri')

    # Disclaimer
    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_p.paragraph_format.space_before = Pt(16)
    _add_run(disc_p, "Disclaimer: ", bold=True, size=9, color=MUTED)
    _add_run(disc_p, "This is an advisory tool. Recommendations are AI-generated and should be verified independently. Always confirm scope, APC, and submission guidelines with the journal directly.", size=9, color=MUTED)

    # ── Write to buffer ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
